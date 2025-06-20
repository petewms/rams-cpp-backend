from flask import Flask, request, send_file
from flask_cors import CORS
from docx import Document
import tempfile
import os
import pandas as pd

app = Flask(__name__)
CORS(app)

@app.route("/generate-cpp", methods=["POST"])
def generate_cpp():
    file = request.files["quote"]
    if not file:
        return "No file uploaded", 400

    try:
        xls = pd.ExcelFile(file)
        df = xls.parse(xls.sheet_names[0], header=None)

        job_number = df.iloc[1, 10] if len(df.columns) > 10 and not pd.isna(df.iloc[1, 10]) else "TBC"
        site_address = df.iloc[10, 1] if len(df) > 10 and not pd.isna(df.iloc[10, 1]) else "Site Address"

        scope_lines = []
        for i in range(15, len(df)):
            desc = df.iloc[i, 2] if len(df.columns) > 2 else None
            if isinstance(desc, str) and desc.strip():
                scope_lines.append(f"- {desc.strip()}")
        scope_text = "\n".join(scope_lines) if scope_lines else "Scope of works"

        template_path = os.path.join("templates", "cpp_template.docx")
        doc = Document(template_path)

        replacements = {
            "{{SiteAddress}}": site_address,
            "{{ScopeOfWorks}}": scope_text,
            "{{Client}}": "LiveWest",
            "{{JobNumber}}": job_number
        }

        for p in doc.paragraphs:
            for key, val in replacements.items():
                if key in p.text:
                    p.text = p.text.replace(key, val)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, val in replacements.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, val)

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        return send_file(tmp.name, as_attachment=True, download_name="CPP_Filled.docx")
    except Exception as e:
        return str(e), 500